
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DeepSeek AI 自动写作脚本
使用 DeepSeek API 根据 prompt.txt 和 example.txt 生成小说并输出为 Word 文档
/Users/jacksu/Desktop/AI-auto-writer-script/.venv/bin/python auto_writer.py
"""

import requests
import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv

# 加载环境变量
load_dotenv()

# DeepSeek API 配置
API_KEY = os.getenv('DEEPSEEK_API_KEY')
BASE_URL = os.getenv('DEEPSEEK_BASE_URL', 'https://api.deepseek.com/v1/chat/completions')

if not API_KEY:
    raise ValueError("请在.env文件中设置DEEPSEEK_API_KEY环境变量")

class DeepSeekAutoWriter:
    def __init__(self):
        self.api_key = API_KEY
        self.base_url = BASE_URL
        self.headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
    
    def clean_markdown_format(self, text):
        """清理文本中的Markdown格式符号，转换为纯文本"""
        import re
        
        # 移除各种Markdown格式符号
        text = re.sub(r'#{1,6}\s*', '', text)  # 移除标题符号 #
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)  # 移除粗体 **text**
        text = re.sub(r'\*(.*?)\*', r'\1', text)  # 移除斜体 *text*
        text = re.sub(r'__(.*?)__', r'\1', text)  # 移除粗体 __text__
        text = re.sub(r'_(.*?)_', r'\1', text)  # 移除斜体 _text_
        text = re.sub(r'^[-*+]\s+', '', text, flags=re.MULTILINE)  # 移除列表符号
        text = re.sub(r'^\d+\.\s+', '', text, flags=re.MULTILINE)  # 移除有序列表
        text = re.sub(r'`(.*?)`', r'\1', text)  # 移除代码符号
        text = re.sub(r'>\s*', '', text, flags=re.MULTILINE)  # 移除引用符号
        text = re.sub(r'\[([^\]]*)\]\([^)]*\)', r'\1', text)  # 移除链接格式
        
        # 移除章节标题和分隔符
        text = re.sub(r'^第[一二三四五六七八九十\d]+章.*?$', '', text, flags=re.MULTILINE)  # 移除章节标题
        text = re.sub(r'^第\d+章.*?$', '', text, flags=re.MULTILINE)  # 移除数字章节标题
        text = re.sub(r'^[-=]{3,}$', '', text, flags=re.MULTILINE)  # 移除分隔线
        text = re.sub(r'^\s*---+\s*$', '', text, flags=re.MULTILINE)  # 移除---分隔符
        text = re.sub(r'^\s*===+\s*$', '', text, flags=re.MULTILINE)  # 移除===分隔符
        
        # 清理多余的空行
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # 将多个空行压缩为双空行
        
        return text.strip()

    def read_file(self, filename):
        """读取文件内容"""
        try:
            with open(filename, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except FileNotFoundError:
            print(f"错误：找不到文件 {filename}")
            return None
        except Exception as e:
            print(f"读取文件 {filename} 时出错：{e}")
            return None
    
    def call_deepseek_api(self, messages, max_tokens=2000, max_retries=5):
        """调用 DeepSeek API，带重试机制和更好的错误处理"""
        import time
        
        payload = {
            "model": "deepseek-chat",
            "messages": messages,
            "max_tokens": max_tokens,
            "temperature": 0.7,
            "top_p": 0.9,
            "stream": False
        }
        
        for attempt in range(max_retries):
            try:
                print(f"正在调用API... (尝试 {attempt + 1}/{max_retries})")
                
                # 创建session以复用连接
                session = requests.Session()
                session.headers.update(self.headers)
                
                response = session.post(
                    self.base_url,
                    json=payload,
                    timeout=(30, 180)  # 连接超时30s，读取超时180s
                )
                
                if response.status_code == 200:
                    result = response.json()
                    session.close()
                    return result
                else:
                    print(f"API 调用失败，状态码：{response.status_code}")
                    print(f"错误信息：{response.text[:200]}...")
                    session.close()
                    
            except requests.exceptions.Timeout as e:
                print(f"请求超时: {str(e)} (尝试 {attempt + 1}/{max_retries})")
            except requests.exceptions.ConnectionError as e:
                print(f"连接错误: {str(e)} (尝试 {attempt + 1}/{max_retries})")
            except requests.exceptions.RequestException as e:
                print(f"请求异常: {str(e)} (尝试 {attempt + 1}/{max_retries})")
            except Exception as e:
                print(f"其他错误: {str(e)} (尝试 {attempt + 1}/{max_retries})")
            
            # 如果不是最后一次尝试，等待后重试
            if attempt < max_retries - 1:
                wait_time = (attempt + 1) * 5  # 递增等待时间
                print(f"等待 {wait_time} 秒后重试...")
                time.sleep(wait_time)
        
        print("所有重试都失败了")
        return None
    
    def generate_novel_title(self, outline):
        """根据大纲生成小说标题"""
        simple_messages = [
            {
                "role": "system",
                "content": "你是专业的小说编辑，擅长根据故事大纲创作吸引人的小说标题。"
            },
            {
                "role": "user",
                "content": f"""根据以下小说大纲，为这部小说创作一个吸引人的标题：

{outline[:1000]}...

要求：
1. 标题要简洁有力，不超过15个字
2. 能体现故事的核心主题
3. 具有吸引力和文学性
4. 适合都市言情小说风格

请只返回标题，不要其他内容。"""
            }
        ]
        
        print("正在生成小说标题...")
        response = self.call_deepseek_api(simple_messages, max_tokens=100)
        
        if response and 'choices' in response:
            title = response['choices'][0]['message']['content'].strip()
            # 清理标题，去掉可能的引号和特殊字符
            title = title.replace('"', '').replace("'", '').replace('《', '').replace('》', '')
            title = title.replace('\n', '').strip()
            return title[:20] if len(title) > 20 else title  # 限制标题长度
        return "AI生成小说"
    
    def generate_novel_outline(self, system_prompt, example_text, user_request):
        """生成小说大纲"""
        messages = [
            {
                "role": "system",
                "content": system_prompt
            },
            {
                "role": "user",
                "content": f"""请根据以下参考文本的风格和结构，{user_request}

参考文本示例：
{example_text[:2000]}...

请按照要求生成一个完整的短篇小说大纲：
1. 总字数控制在1万字左右
2. 分为3章，每章约3000-3500字
3. 确保故事有完整的开头、发展、高潮、结局
4. 先生成一句话概括，然后生成详细的3章小说提纲
5. 请使用纯文本格式输出，不要使用任何Markdown格式符号（如#、*、-等）"""
            }
        ]
        
        print("正在生成小说大纲...")
        response = self.call_deepseek_api(messages)
        
        if response and 'choices' in response:
            return response['choices'][0]['message']['content']
        return None
    
    def generate_chapter(self, system_prompt, example_text, outline, chapter_num, total_chapters=3, previous_chapters=""):
        """生成指定章节"""
        # 根据总章节数调整每章字数，确保总字数约一万字
        target_words = 16000 // total_chapters
        
        simple_system = f"""你是专业小说作家。要求：
1. 语言通俗易懂，口语化
2. 情节紧凑有张力
3. 人物性格鲜明
4. 这是第{chapter_num}章，共{total_chapters}章
5. 内容约{target_words}字，确保情节完整
6. 使用纯文本格式，不要使用Markdown格式符号"""
        
        simple_example = example_text[:500] if len(example_text) > 500 else example_text
        simple_outline = outline[:600] if len(outline) > 600 else outline
        
        # 根据章节位置给出不同的创作指导
        if chapter_num == 1:
            chapter_guide = "开篇要有强烈冲突，设置悬念，引出主要人物和矛盾"
        elif chapter_num == total_chapters:
            chapter_guide = "结尾章节，要解决所有冲突，给出完整结局，让故事有圆满收尾"
        else:
            chapter_guide = "中间章节，推进情节发展，加深人物关系，为结局做铺垫"
        
        messages = [
            {
                "role": "system",
                "content": simple_system
            },
            {
                "role": "user",
                "content": f"""根据大纲写第{chapter_num}章（共{total_chapters}章），约{target_words}字：

大纲：{simple_outline}

参考风格：{simple_example}

创作要求：{chapter_guide}

请写第{chapter_num}章完整内容，约{target_words}字。
重要要求：
1. 使用纯文本格式，不要使用任何Markdown格式符号（如#、*、**、-等）
2. 不要写章节标题（如"第一章"、"第二章"等）
3. 不要使用分隔符（如---、===等）
4. 直接开始写故事正文内容"""
            }
        ]
        
        print(f"正在生成第{chapter_num}章（共{total_chapters}章，目标{target_words}字）...")
        return self.call_deepseek_api(messages, max_tokens=int(target_words * 1.5))
    
    def create_word_document(self, title, content):
        """创建 Word 文档"""
        doc = Document()
        
        # 设置文档标题 - 3号字体（16pt）
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Pt(16)  # 3号字体 = 16pt
        title_run.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加空行
        doc.add_paragraph()
        
        # 添加生成时间 - 小字体
        time_info = f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}"
        time_paragraph = doc.add_paragraph(time_info)
        time_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # 设置时间字体为小五号（9pt）
        for run in time_paragraph.runs:
            run.font.size = Pt(9)
        
        # 添加分隔线
        separator_paragraph = doc.add_paragraph("=" * 50)
        separator_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 设置分隔线字体
        for run in separator_paragraph.runs:
            run.font.size = Pt(9)
        
        # 添加正文内容 - 5号字体（10.5pt）
        paragraphs = content.split('\n')
        for paragraph_text in paragraphs:
            if paragraph_text.strip():
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(paragraph_text.strip())
                # 设置正文字体为5号（10.5pt）
                run.font.size = Pt(10.5)
        
        return doc
    
    def save_document(self, doc, filename):
        """保存文档"""
        try:
            doc.save(filename)
            print(f"文档已保存为：{filename}")
            return True
        except Exception as e:
            print(f"保存文档时出错：{e}")
            return False
    
    def run(self, user_request="生成一部现代都市言情小说"):
        """主运行函数"""
        print("=== DeepSeek AI 自动写作脚本 ===")
        
        # 读取提示词文件
        print("读取系统提示词...")
        system_prompt = self.read_file("prompt.txt")
        if not system_prompt:
            return False
        
        # 读取示例文本
        print("读取参考文本...")
        example_text = self.read_file("example.txt")
        if not example_text:
            return False
        
        # 生成小说大纲
        outline = self.generate_novel_outline(system_prompt, example_text, user_request)
        if not outline:
            print("生成大纲失败")
            return False
        
        # 清理大纲中的Markdown格式
        clean_outline = self.clean_markdown_format(outline)
        
        print("\n=== 生成的小说大纲 ===")
        print(clean_outline)
        print("\n" + "="*50)
        
        # 根据大纲生成标题
        novel_title = self.generate_novel_title(clean_outline)
        print(f"\n=== 生成的小说标题 ===")
        print(f"《{novel_title}》")
        print("\n" + "="*50)
        
        # 生成小说内容
        novel_content = f"小说大纲：\n{clean_outline}\n\n"
        novel_content += "="*50 + "\n正文内容：\n" + "="*50 + "\n\n"
        
        previous_chapters = ""
        total_chapters = 3  # 改为3章
        
        # 生成3章内容，每章约3000-3500字
        for chapter_num in range(1, total_chapters + 1):
            response = self.generate_chapter(
                system_prompt, 
                example_text, 
                clean_outline, 
                chapter_num,
                total_chapters, 
                previous_chapters
            )
            
            if response and 'choices' in response:
                raw_content = response['choices'][0]['message']['content']
                # 清理Markdown格式，确保是纯文本
                chapter_content = self.clean_markdown_format(raw_content)
                # 直接添加内容，不添加章节标题
                novel_content += f"{chapter_content}\n\n"
                previous_chapters += chapter_content + "\n\n"
                print(f"第{chapter_num}章生成完成")
            else:
                print(f"第{chapter_num}章生成失败，跳过")
                # 继续生成下一章而不是终止整个过程
        
        # 创建并保存 Word 文档
        print("正在创建 Word 文档...")
        doc = self.create_word_document(novel_title, novel_content)
        
        # 生成文件名，清理标题中的特殊字符
        safe_title = novel_title.replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{safe_title}_{timestamp}.docx"
        
        if self.save_document(doc, filename):
            print(f"\n=== 生成完成 ===")
            print(f"小说标题：{novel_title}")
            print(f"输出文件：{filename}")
            print(f"总字数：约{len(novel_content)}字")
            print(f"目标：一万字左右的完整短篇小说")
            return True
        else:
            return False

def main():
    """主函数"""
    writer = DeepSeekAutoWriter()
    
    # 用户可以在这里自定义请求
    user_request = input("请输入您想要生成的小说类型和要求（直接回车使用默认设置）：").strip()
    if not user_request:
        user_request = "生成一部一万字左右的现代都市言情短篇小说，要有完整的故事情节，包含爱情、成长、励志元素"
    
    success = writer.run(user_request)
    
    if success:
        print("\n小说生成成功！")
    else:
        print("\n小说生成失败，请检查错误信息并重试。")

if __name__ == "__main__":
    main()