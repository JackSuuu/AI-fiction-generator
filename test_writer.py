#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DeepSeek AI 自动写作脚本 - 测试版本
"""

import requests
import json
import os
from datetime import datetime
from docx import Document
from dotenv import load_dotenv

# 加载环境变量
load_dotenv()

# DeepSeek API 配置
API_KEY = os.getenv('DEEPSEEK_API_KEY')
BASE_URL = os.getenv('DEEPSEEK_BASE_URL', 'https://api.deepseek.com/v1/chat/completions')

if not API_KEY:
    raise ValueError("请在.env文件中设置DEEPSEEK_API_KEY环境变量")

def call_deepseek_api(messages, max_tokens=2000):
    """调用 DeepSeek API"""
    headers = {
        "Authorization": f"Bearer {API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "deepseek-chat",
        "messages": messages,
        "max_tokens": max_tokens,
        "temperature": 0.7,
        "stream": False
    }
    
    try:
        print("正在调用API...")
        response = requests.post(
            BASE_URL,
            headers=headers,
            json=payload,
            timeout=90
        )
        
        if response.status_code == 200:
            return response.json()
        else:
            print(f"API 调用失败，状态码：{response.status_code}")
            print(f"错误信息：{response.text}")
            return None
            
    except Exception as e:
        print(f"API 调用出错：{e}")
        return None

def read_file(filename):
    """读取文件内容"""
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            return f.read().strip()
    except Exception as e:
        print(f"读取文件 {filename} 时出错：{e}")
        return None

def generate_chapter():
    """生成测试章节"""
    # 读取提示词和示例
    prompt = read_file("prompt.txt")
    example = read_file("example.txt")
    
    if not prompt or not example:
        return None
    
    # 简化prompt
    simplified_prompt = """你是一位专业的网络小说作家，擅长写都市言情小说。
    要求：
    1. 语言通俗易懂，口语化表达
    2. 情节紧凑，有冲突有张力
    3. 人物性格鲜明
    4. 内容至少1500字
    """
    
    # 简化示例
    simplified_example = example[:1000] if len(example) > 1000 else example
    
    messages = [
        {
            "role": "system",
            "content": simplified_prompt
        },
        {
            "role": "user",
            "content": f"""参考以下文本风格，写一个都市言情小说的第一章，主角是一个遭遇背叛的都市女性：

参考风格：
{simplified_example}

要求：
1. 开篇要有冲击力
2. 至少1500字
3. 情节要有起伏
4. 人物要有个性"""
        }
    ]
    
    response = call_deepseek_api(messages, max_tokens=2500)
    
    if response and 'choices' in response:
        return response['choices'][0]['message']['content']
    return None

def save_to_word(content, filename):
    """保存到Word文档"""
    try:
        doc = Document()
        
        # 添加标题
        title = doc.add_paragraph()
        title.add_run("AI生成小说测试章节").bold = True
        
        # 添加时间
        doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("=" * 50)
        
        # 添加内容
        paragraphs = content.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        doc.save(filename)
        print(f"文档已保存为：{filename}")
        return True
    except Exception as e:
        print(f"保存失败：{e}")
        return False

def main():
    print("=== DeepSeek AI 写作测试 ===")
    
    content = generate_chapter()
    if content:
        print("\n生成成功！内容预览：")
        print("-" * 50)
        print(content[:300] + "..." if len(content) > 300 else content)
        print("-" * 50)
        
        filename = f"test_novel_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        save_to_word(content, filename)
    else:
        print("生成失败")

if __name__ == "__main__":
    main()
